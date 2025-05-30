import streamlit as st
import os
from docx import Document
import logging
from typing import List
from app_with_complete_gemini import RFPGenerator
import google.generativeai as genai

# Configure logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Streamlit app
st.title("RFP Generator Dashboard")

# Initialize session state for file uploads and outputs
if "rfp_files" not in st.session_state:
    st.session_state["rfp_files"] = []
if "proposal_files" not in st.session_state:
    st.session_state["proposal_files"] = []
if "generated_files" not in st.session_state:
    st.session_state["generated_files"] = []

# Sidebar for file uploads
st.sidebar.header("Upload RFP and Proposal Files")
rfp_upload = st.sidebar.file_uploader("Upload RFP Files (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)
proposal_upload = st.sidebar.file_uploader("Upload Proposal Files (PDF/DOCX/TXT)", type=["pdf", "docx", "txt"], accept_multiple_files=True)

# Save uploaded files
def save_uploaded_files(uploaded_files, folder: str, prefix="uploaded") -> List[str]:
    """Save uploaded files to a folder and return file paths"""
    folder = os.path.abspath(folder)
    os.makedirs(folder, exist_ok=True)
    file_paths = []
    for uploaded_file in uploaded_files:
        file_path = os.path.join(folder, f"{prefix}_{uploaded_file.name}")
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        file_paths.append(file_path)
    return file_paths

# Process uploaded files
if rfp_upload:
    st.session_state["rfp_files"] = save_uploaded_files(rfp_upload, "uploads/rfp", "rfp")
    rfp_count = len(st.session_state["rfp_files"])
    st.sidebar.write(f"Uploaded {rfp_count} RFP file{'s' if rfp_count != 1 else ''}")
if proposal_upload:
    st.session_state["proposal_files"] = save_uploaded_files(proposal_upload, "uploads/proposal", "proposal")
    proposal_count = len(st.session_state["proposal_files"])
    st.sidebar.write(f"Uploaded {proposal_count} Proposal file{'s' if proposal_count != 1 else ''}")

# Button to generate RFP
if st.button("Generate RFP"):
    if not st.session_state["rfp_files"] or not st.session_state["proposal_files"]:
        st.error("Please upload at least one RFP file and one proposal file.")
    else:
        try:
            # Initialize backend
            genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))
            generator = RFPGenerator(gemini_api_key=os.getenv("GOOGLE_API_KEY"))

            # Generate RFP
            with st.spinner("Generating RFP..."):
                result = generator.generate_rfp(
                    rfp_files=st.session_state["rfp_files"],
                    proposal_files=st.session_state["proposal_files"],
                    output_dir="generated_files"
                )
                generated_files = [os.path.abspath(f) for f in result.get("files_created", [])]
                st.session_state["generated_files"] = generated_files
            st.success(f"Generated {len(generated_files)} files!")

        except Exception as e:
            st.error(f"Error generating RFP: {str(e)}")
            logger.error(f"Error generating RFP: {e}")

# Display generated DOCX files in main section
if st.session_state["generated_files"]:
    st.header("Generated RFP Output")
    for file_path in st.session_state["generated_files"]:
        file_name = os.path.basename(file_path)
        if file_name.endswith("validation_report.txt"):
            continue
        absolute_path = os.path.abspath(file_path)
        logger.debug(f"Attempting to access file: {absolute_path}")
        
        if not os.path.exists(absolute_path):
            st.warning(f"File not found: {file_name} at {absolute_path}")
            logger.error(f"File not found: {absolute_path}")
            continue

        if file_path.endswith(".docx"):
            try:
                doc = Document(absolute_path)
                content = "\n".join([para.text for para in doc.paragraphs[:100]])[:1000]
                with st.expander(f"Preview: {file_name}"):
                    st.text_area("Content Preview", content, height=200)
            except Exception as e:
                st.warning(f"Could not preview {file_name}: {e}")
                logger.error(f"Preview error for {file_name}: {e}")
        elif file_path.endswith(".txt"):
            try:
                with open(absolute_path, 'r', encoding='utf-8') as f:
                    content = f.read()[:1000]
                with st.expander(f"Preview: {file_name}"):
                    st.text_area("Content Preview", content, height=200)
            except Exception as e:
                st.warning(f"Could not preview {file_name}: {e}")
                logger.error(f"Preview error for {file_name}: {e}")

        # Download button
        try:
            with open(absolute_path, "rb") as f:
                st.download_button(
                    label=f"Download {file_name}",
                    data=f,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_path.endswith(".docx") else "text/plain"
                )
        except Exception as e:
            st.warning(f"Could not provide download for {file_name}: {e}")
            logger.error(f"Download error for {file_name}: {e}")

# Sidebar dashboard for all generated files (including validation report)
st.sidebar.header("Output Dashboard")
if st.session_state["generated_files"]:
    for file_path in st.session_state["generated_files"]:
        file_name = os.path.basename(file_path)
        absolute_path = os.path.abspath(file_path)
        logger.debug(f"Sidebar: Attempting to access file: {absolute_path}")

        if not os.path.exists(absolute_path):
            st.sidebar.warning(f"File not found: {file_name} at {absolute_path}")
            logger.error(f"Sidebar: File not found: {absolute_path}")
            continue

        st.sidebar.write(f"**{file_name}**")
        # Preview in sidebar
        if file_path.endswith(".docx"):
            try:
                doc = Document(absolute_path)
                content = "\n".join([para.text for para in doc.paragraphs[:5]])[:500]  # Shorter preview for sidebar
                with st.sidebar.expander(f"Preview: {file_name}"):
                    st.text_area("Content Preview", content, height=150, key=f"sidebar_preview_{file_name}")
            except Exception as e:
                st.sidebar.warning(f"Could not preview {file_name}: {e}")
                logger.error(f"Sidebar: Preview error for {file_name}: {e}")
        elif file_path.endswith(".txt"):
            try:
                with open(absolute_path, 'r', encoding='utf-8') as f:
                    content = f.read()[:500]
                with st.sidebar.expander(f"Preview: {file_name}"):
                    st.text_area("Content Preview", content, height=150, key=f"sidebar_preview_{file_name}")
            except Exception as e:
                st.sidebar.warning(f"Could not preview {file_name}: {e}")
                logger.error(f"Sidebar: Preview error for {file_name}: {e}")

        # Download button in sidebar
        try:
            with open(absolute_path, "rb") as f:
                st.sidebar.download_button(
                    label=f"Download {file_name}",
                    data=f,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if file_path.endswith(".docx") else "text/plain",
                    key=f"sidebar_download_{file_name}"
                )
        except Exception as e:
            st.sidebar.warning(f"Could not provide download for {file_name}: {e}")
            logger.error(f"Sidebar: Download error for {file_name}: {e}")
else:
    st.sidebar.write("No output files generated yet.")