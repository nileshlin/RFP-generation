import os
import glob
from pathlib import Path
import re
import json
from typing import List, Dict, Optional
from dataclasses import dataclass, asdict
from datetime import datetime, timedelta
import logging
import pdfplumber
import spacy
# import cohere
import google.generativeai as genai
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
from spacy.cli import download

load_dotenv()

# Data classes
@dataclass
class RFPRequirement:
    section_title: str
    content: str
    requirement_type: str
    priority: float = 0.5
    mandatory: bool = True

@dataclass
class ProposalSection:
    section_title: str
    content: str
    section_type: str
    source_file: str
    confidence_score: float = 0.0

class DocumentExtractor:
    """Extract text from PDF and DOCX files"""
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def extract_text(self, file_path: str) -> str:
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                return self._extract_pdf(file_path)
            elif ext == '.docx':
                return self._extract_docx(file_path)
            else:
                self.logger.error(f"Unsupported file format: {file_path}")
                return ""
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return ""

    def _extract_pdf(self, pdf_path: str) -> str:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text

    def _extract_docx(self, docx_path: str) -> str:
        doc = Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        return text

class RFPParser:
    """Parse RFP documents to extract requirements and config"""
    def __init__(self):
        # self.nlp = spacy.load("en_core_web_sm")
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            download("en_core_web_sm")
            self.nlp = spacy.load("en_core_web_sm")

        self.common_patterns = [
            r"Scope\s+of\s+Work",
            r"Submission\s+Requirements",
            r"Staffing\s+Requirements",
            r"Evaluation\s+Criteria",
            r"Qualifications",
            r"Insurance\s+Requirements",
            r"Terms\s+and\s+Conditions",
            r"Mandatory\s+Forms",
            r"Contract\s+Terms",
            r"Data\s+Privacy",
            r"Bidder\s+Checklist",
            r"Compliance\s+Requirements",
            r"Payment\s+Terms",
            r"Service\s+Specifications",
            r"Reporting\s+Requirements"
        ]
        self.priority_keywords = [
            r"mandatory|required|must|shall",
            r"compliance|regulation|law",
            r"submission|deadline|proposal",
            r"qualification|credential|license",
            r"insurance|liability",
            r"contract|agreement",
            r"safety|health"
        ]
        self.logger = logging.getLogger(__name__)

    def parse_requirements(self, rfp_text: str) -> List[RFPRequirement]:
        """Parse all RFP requirements, ranking by priority"""
        requirements = []
        doc = self.nlp(rfp_text)
        current_section = None
        current_content = []
        section_count = 0

        for sent in doc.sents:
            sent_text = sent.text.strip()
            is_section = False
            for pattern in self.common_patterns:
                if re.search(pattern, sent_text, re.IGNORECASE):
                    if current_section and current_content:
                        priority = self._calculate_priority(" ".join(current_content))
                        requirements.append(RFPRequirement(
                            section_title=current_section,
                            content=" ".join(current_content),
                            requirement_type=self._infer_requirement_type(current_section),
                            priority=priority
                        ))
                        section_count += 1
                    current_section = sent_text[:100]
                    current_content = []
                    is_section = True
                    break
            if not is_section and current_section:
                current_content.append(sent_text)

        if current_section and current_content:
            priority = self._calculate_priority(" ".join(current_content))
            requirements.append(RFPRequirement(
                section_title=current_section,
                content=" ".join(current_content),
                requirement_type=self._infer_requirement_type(current_section),
                priority=priority
            ))

        # Sort by priority (descending)
        requirements.sort(key=lambda x: x.priority, reverse=True)
        self.logger.info(f"Parsed {len(requirements)} requirements")
        return requirements
    
    def _calculate_priority(self, content: str) -> float:
        """Assign priority based on keywords"""
        priority = 0.5
        for pattern in self.priority_keywords:
            if re.search(pattern, content, re.IGNORECASE):
                priority += 0.1
        return min(priority, 1.0)
    
    def _infer_requirement_type(self, section_title: str) -> str:
        section_title = section_title.lower()
        if 'scope' in section_title:
            return 'scope'
        elif 'submission' in section_title:
            return 'submission'
        elif 'staff' in section_title or 'qualification' in section_title:
            return 'staffing'
        elif 'insurance' in section_title:
            return 'insurance'
        elif 'terms' in section_title or 'condition' in section_title:
            return 'terms'
        elif 'forms' in section_title:
            return 'forms'
        elif 'contract' in section_title:
            return 'contract'
        elif 'privacy' in section_title or 'security' in section_title:
            return 'privacy'
        elif 'checklist' in section_title:
            return 'checklist'
        elif 'compliance' in section_title:
            return 'compliance'
        elif 'payment' in section_title:
            return 'payment'
        elif 'service' in section_title:
            return 'service'
        elif 'reporting' in section_title:
            return 'reporting'
        return 'general'

    def extract_config(self, rfp_text: str, proposal_texts: List[str], gemini_client) -> Dict:
        """Generate config dynamically using Gemini with regex/spaCy fallback"""
        self.logger.info("Generating config using Gemini")
        prompt = f"""You are extracting configuration details for a Request for Proposals (RFP) for a school district.
        Input RFP Text:
        {rfp_text[:1000]}

        Input Proposal Texts:
        {''.join([text[:500] for text in proposal_texts])}

        Task: Extract or infer the following configuration fields based on the input texts. If a field cannot be extracted, provide a reasonable default value relevant to a school district issuing an RFP for nursing services for the 2025/2026 school year. Return the result as a JSON object.

        Fields:
        - district_name: Name of the school district (e.g., "Toms River Regional School District")
        - rfp_number: Unique identifier for the RFP (e.g., "2025/2026")
        - district_address: Full address of the district (e.g., "123 Main St, Toms River, NJ 08753")
        - district_phone: Contact phone number (e.g., "(732) 555-1234")
        - district_fax: Fax number (e.g., "(732) 555-5678")
        - district_website: Website URL (e.g., "www.tomsriverschools.org")
        - service_type: Type of service (e.g., "Nursing Services")
        - submission_deadline: Proposal submission deadline (e.g., "June 26, 2025, 10:00 AM")
        - issue_date: RFP issue date (e.g., "May 27, 2025")
        - contact_name: Contact person’s name (e.g., "Jane Doe")
        - contact_email: Contact email (e.g., "purchasing@tomsriverschools.org")
        - school_year: Academic year (e.g., "2025/2026")
        - student_count: Number of students (e.g., "10,000")
        - school_count: Number of schools (e.g., "12")

        Response Format:
        ```json
        {{
        "district_name": "...",
        "rfp_number": "...",
        ...
        }}
        ```
        """
        config = {}
        try:
            response = gemini_client.generate_content(
                prompt,
                generation_config={
                    "max_output_tokens": 250,
                    "temperature": 0.4,
                    "stop_sequences": ["```"]
                }
            )
            config_text = response.text.strip()
            self.logger.debug(f"Gemini response: {config_text}")
            if not config_text:
                raise ValueError("Empty response from Gemini")
            if config_text.startswith("```json"):
                config_text = config_text[7:]
            if config_text.endswith("```"):
                config_text = config_text[:-3]
            config = json.loads(config_text)
            self.logger.debug(f"Generated config: {json.dumps(config, indent=2)}")
        except (ValueError, json.JSONDecodeError, Exception) as e:
            self.logger.error(f"Error generating config with Gemini: {e}")
            config = {}

        if not config:
            config = {
                'district_name': 'Toms River Regional School District',
                'rfp_number': '2025/2026',
                'district_address': '1234 Main Street, Toms River, NJ 08753',
                'district_phone': '(732) 555-1234',
                'district_fax': '(732) 555-5678',
                'district_website': 'www.tomsriverschools.org',
                'service_type': 'Nursing Services',
                'submission_deadline': (datetime.now() + timedelta(days=30)).strftime('%B %d, %Y, 10:00 AM'),
                'issue_date': datetime.now().strftime('%B %d, %Y'),
                'contact_name': 'Jane Doe',
                'contact_email': 'purchasing@tomsriverschools.org',
                'school_year': '2025/2026',
                'student_count': '10,000',
                'school_count': '12'
            }
            try:
                doc = self.nlp(rfp_text) if rfp_text else None
                if doc:
                    address_pattern = r'(\d+\s+[A-Za-z\s]+,\s+[A-Z]{2}\s+\d{5})'
                    email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
                    phone_pattern = r'\(\d{3}\)\s*\d{3}-\d{4}'
                    date_pattern = r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}'
                    district_pattern = r'([A-Za-z\s]+(?:School District|Regional School District|County School District))'
                    for sent in doc.sents:
                        match = re.search(district_pattern, sent.text, re.IGNORECASE)
                        if match:
                            district_text = match.group(1).strip()
                            if len(district_text.split()) <= 6 and 'school district' in district_text.lower():
                                config['district_name'] = district_text
                        match = re.search(address_pattern, sent.text)
                        if match:
                            config['district_address'] = re.sub(r'\n', ' ', match.group(1)).strip()
                        if re.search(email_pattern, sent.text):
                            config['contact_email'] = re.search(email_pattern, sent.text).group()
                        if re.search(phone_pattern, sent.text):
                            config['district_phone'] = re.search(phone_pattern, sent.text).group()
                        if 'deadline' in sent.text.lower():
                            match = re.search(date_pattern, sent.text)
                            if match:
                                config['submission_deadline'] = f"{match.group()}, 10:00 AM"
                        if 'service' in sent.text.lower():
                            for ent in sent.ents:
                                if ent.label_ in ['ORG', 'PRODUCT'] and 'district' not in ent.text.lower():
                                    config['service_type'] = ent.text
                for proposal_text in proposal_texts:
                    doc = self.nlp(proposal_text)
                    for sent in doc.sents:
                        match = re.search(district_pattern, sent.text, re.IGNORECASE)
                        if match:
                            district_text = match.group(1).strip()
                            if len(district_text.split()) <= 6 and 'school district' in district_text.lower():
                                config['district_name'] = district_text
            except Exception as e:
                self.logger.error(f"Error in fallback extraction: {e}")

        return config
    
    def extract_dynamic_sections(self, rfp_text: str, proposal_texts: List[str]) -> Dict:
        """Extract information for dynamic RFP sections using regex and spaCy"""
        dynamic_sections = {
            'general_terms_and_conditions': {},
            # 'background': {},
            'scope_of_services': {},
            'timeline': {},
            # 'insurance_requirements': {},
            'terms_and_conditions': {},
            'attachments': {},
            'appendices': {},
            'submitter_information': {}
        }
        try:
            doc = self.nlp(rfp_text) if rfp_text else None
            if doc:
                # Patterns for extraction
                insurance_pattern = r'(?:insurance|liability|coverage)\s*(?:requirements|policy|minimum)\s*[^.]*\b(?:general liability|professional liability|workers[\'\s]*compensation|automobile liability)\b[^.]*'
                terms_pattern = r'(?:terms\s*(?:and|&)\s*conditions|contract\s*(?:award|rejection|payment|compliance|law|responsibilities))\s*[^.]*'
                forms_pattern = r'(?:mandatory\s*forms|required\s*documents|submission\s*forms)\s*[^.]*\b(?:affirmative\s*action|non-collusion|political\s*contribution|ownership\s*disclosure)\b[^.]*'
                appendix_pattern = r'(?:appendix|appendices|checklist|bidder\s*information)\s*[^.]*'
                submitter_pattern = r'(?:submitter|bidder|vendor)\s*(?:information|details|form)\s*[^.]*'

                for sent in doc.sents:
                    text = sent.text.lower()
                    # Introduction
                    if 'introduction' in text or 'purpose' in text or 'background' in text or 'district' in text or 'overview' in text:
                        dynamic_sections['general_terms_and_conditions']['text'] = sent.text.strip()
                    # Background
                    # if 'background' in text or 'district' in text or 'overview' in text:
                    #     dynamic_sections['background']['text'] = sent.text.strip()
                    # Objective
                    if 'objective' in text or 'goal' in text or 'aim' in text:
                        dynamic_sections['scope_of_services']['text'] = sent.text.strip()
                    # Timeline
                    if 'timeline' in text or 'schedule' in text or 'deadline' in text:
                        dynamic_sections['timeline']['text'] = sent.text.strip()
                    # Insurance Requirements
                    # if re.search(insurance_pattern, text, re.IGNORECASE):
                    #     insurance_items = dynamic_sections['insurance_requirements'].get('items', [])
                    #     insurance_items.append(sent.text.strip())
                    #     dynamic_sections['insurance_requirements']['items'] = insurance_items
                    # Terms and Conditions
                    if re.search(terms_pattern, text, re.IGNORECASE):
                        terms_items = dynamic_sections['terms_and_conditions'].get('items', [])
                        terms_items.append(sent.text.strip())
                        dynamic_sections['terms_and_conditions']['items'] = terms_items
                    # Mandatory Forms
                    if re.search(forms_pattern, text, re.IGNORECASE):
                        forms = dynamic_sections['attachments'].get('forms', [])
                        form_name = re.search(r'\b(?:affirmative\s*action|non-collusion|political\s*contribution|ownership\s*disclosure)\b', text, re.IGNORECASE)
                        if form_name:
                            forms.append({'name': form_name.group(0).title(), 'fields': [sent.text.strip()]})
                        dynamic_sections['mandatory_forms']['forms'] = forms
                    # Appendices
                    if re.search(appendix_pattern, text, re.IGNORECASE):
                        appendices = dynamic_sections['appendices'].get('items', [])
                        appendices.append(sent.text.strip())
                        dynamic_sections['appendices']['items'] = appendices
                    # Submitter Information
                    if re.search(submitter_pattern, text, re.IGNORECASE):
                        fields = dynamic_sections['submitter_information'].get('fields', [])
                        fields.append(sent.text.strip())
                        dynamic_sections['submitter_information']['fields'] = fields

            # Process proposals
            for proposal_text in proposal_texts:
                doc = self.nlp(proposal_text)
                for sent in doc.sents:
                    text = sent.text.lower()
                    if re.search(insurance_pattern, text, re.IGNORECASE):
                        insurance_items = dynamic_sections['insurance_requirements'].get('items', [])
                        insurance_items.append(sent.text.strip())
                        dynamic_sections['insurance_requirements']['items'] = insurance_items
                    if re.search(terms_pattern, text, re.IGNORECASE):
                        terms_items = dynamic_sections['terms_and_conditions'].get('items', [])
                        terms_items.append(sent.text.strip())
                        dynamic_sections['terms_and_conditions']['items'] = terms_items
                    if re.search(forms_pattern, text, re.IGNORECASE):
                        forms = dynamic_sections['mandatory_forms'].get('forms', [])
                        form_name = re.search(r'\b(?:affirmative\s*action|non-collusion|political\s*contribution|ownership\s*disclosure)\b', text, re.IGNORECASE)
                        if form_name:
                            forms.append({'name': form_name.group(0).title(), 'fields': [sent.text.strip()]})
                        dynamic_sections['mandatory_forms']['forms'] = forms
                    if re.search(appendix_pattern, text, re.IGNORECASE):
                        appendices = dynamic_sections['appendices'].get('items', [])
                        appendices.append(sent.text.strip())
                        dynamic_sections['appendices']['items'] = appendices
                    if re.search(submitter_pattern, text, re.IGNORECASE):
                        fields = dynamic_sections['submitter_information'].get('fields', [])
                        fields.append(sent.text.strip())
                        dynamic_sections['submitter_information']['fields'] = fields

        except Exception as e:
            self.logger.error(f"Error extracting dynamic sections: {e}")
            # Default values if extraction fails
            dynamic_sections['insurance_requirements']['items'] = [
                "General Liability Insurance: $1,000,000 per occurrence",
                "Professional Liability Insurance: $2,000,000 aggregate",
                "Workers’ Compensation Insurance: Per NJ law",
                "Automobile Liability Insurance: $500,000"
            ]
            dynamic_sections['terms_and_conditions']['items'] = [
                "Award of Contract: Per N.J.S.A. 19:44A-20.7",
                "Rejection of Proposals: District reserves right to reject",
                "Payment Terms: Within 30 days of invoice",
                "Applicable Law: Governed by New Jersey law"
            ]
            dynamic_sections['mandatory_forms']['forms'] = [
                {'name': 'Affirmative Action', 'fields': ['Form AA302 or equivalent']},
                {'name': 'Non-Collusion Affidavit', 'fields': ['Signed affidavit']},
                {'name': 'Political Contribution', 'fields': ['Disclosure form']}
            ]
            dynamic_sections['appendices']['items'] = [
                "Bidder Checklist: Cover Letter, Affirmative Action, etc.",
                "Submitter Information Form"
            ]
            dynamic_sections['submitter_information']['fields'] = [
                "Submitter Name",
                "Signature",
                "Address",
                "Phone",
                "Email"
            ]

        return dynamic_sections

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        current_section = "Introduction"
        current_content = []
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if self._is_section_header(line):
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line
                current_content = []
            else:
                current_content.append(line)
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        return sections

    def _is_section_header(self, line: str) -> bool:
        header_patterns = [
            r'^\d+\.\s+[A-Z]',
            r'^[A-Z][a-z\s]+Requirements?\s*:?$',
            r'^Section\s+\d+$'
        ]
        return any(re.match(pattern, line, re.IGNORECASE) for pattern in header_patterns)

    def _classify_requirement_type(self, title: str) -> str:
        title_lower = title.lower()
        if 'qualification' in title_lower:
            return 'qualification'
        elif 'scope' in title_lower:
            return 'scope'
        elif 'submission' in title_lower:
            return 'submission'
        elif 'evaluation' in title_lower:
            return 'evaluation'
        elif 'staff' in title_lower:
            return 'staffing'
        elif 'reference' in title_lower:
            return 'reference'
        return 'general'

    def _is_mandatory_requirement(self, content: str) -> bool:
        mandatory_keywords = ['must', 'required', 'mandatory', 'shall']
        return any(keyword in content.lower() for keyword in mandatory_keywords)

class ProposalAnalyzer:
    """Analyze past proposals to extract reusable sections"""
    def __init__(self):
        self.nlp = spacy.load("en_core_web_sm")
        self.section_headers = [
            "Company Overview",
            "Staff Qualifications",
            "Service Capability",
            "Experience",
            "Executive Summary",
            "References"
        ]
        self.logger = logging.getLogger(__name__)

    def analyze_proposals(self, proposal_texts: List[str], filenames: List[str]) -> List[ProposalSection]:
        all_sections = []
        for text, filename in zip(proposal_texts, filenames):
            sections = self._extract_sections(text, filename)
            all_sections.extend(self._filter_sensitive_content(sections))
        return all_sections

    def _extract_sections(self, text: str, filename: str) -> List[ProposalSection]:
        sections = []
        section_dict = self._split_into_sections(text)
        for title, content in section_dict.items():
            section_type = self._classify_section_type(title)
            sections.append(ProposalSection(
                section_title=title,
                content=content,
                section_type=section_type,
                source_file=filename
            ))
        return sections

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        sections = {}
        current_section = "Executive Summary"
        current_content = []
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if self._is_section_header(line):
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line
                current_content = []
            else:
                current_content.append(line)
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        return sections

    def _is_section_header(self, line: str) -> bool:
        header_patterns = [
            r'^[A-Z\s]+:$',
            r'^\d+\.\s+[A-Z]',
            r'^[A-Z][a-z\s]+:$',
            r'^Section\s+\d+'
        ]
        return any(re.match(pattern, line, re.IGNORECASE) for pattern in header_patterns)

    def _classify_section_type(self, title: str) -> str:
        title_lower = title.lower()
        if any(word in title_lower for word in ['executive', 'summary']):
            return 'executive_summary'
        elif any(word in title_lower for word in ['company', 'organization', 'background']):
            return 'company_background'
        elif any(word in title_lower for word in ['staff', 'personnel', 'qualification']):
            return 'staff_qualifications'
        elif any(word in title_lower for word in ['experience', 'past', 'history']):
            return 'experience'
        elif any(word in title_lower for word in ['reference', 'client']):
            return 'references'
        elif any(word in title_lower for word in ['capability', 'ability', 'service']):
            return 'service_capability'
        elif any(word in title_lower for word in ['financial', 'rate', 'cost', 'budget']):
            return 'financial'
        return 'general'

    def _filter_sensitive_content(self, sections: List[ProposalSection]) -> List[ProposalSection]:
        financial_patterns = [
            r'\$[\d,]+',
            r'hourly.*rate',
            r'cost.*$',
            r'budget.*\d+$',
            r'price.*\d+$'
        ]
        filtered_sections = []
        for section in sections:
            is_sensitive = section.section_type == 'financial' or any(
                re.search(pattern, section.content.lower()) for pattern in financial_patterns
            )
            if not is_sensitive:
                filtered_sections.append(section)
            else:
                self.logger.info(f"Filtered sensitive section: {section.section_title}")
        return filtered_sections

class ContentMatcher:
    """Match RFP requirements to proposal sections using Gemini"""
    def __init__(self, gemini_api_key: str):
        genai.configure(api_key=gemini_api_key)
        self.model = genai.GenerativeModel("gemini-1.5-flash")
        self.embedding_model = "models/text-embedding-004"
        self.logger = logging.getLogger(__name__)

    def match_requirements(self, requirements: List[RFPRequirement], sections: List[ProposalSection]) -> Dict[str, List[ProposalSection]]:
        self.logger.debug(f"Matching {len(requirements)} requirements with {len(sections)} sections")
        rfp_texts = [f"{req.section_title}: {req.content}" for req in requirements]
        section_texts = [f"{section.section_title}: {section.content}" for section in sections]
        rfp_embeddings = self._get_embeddings(rfp_texts)
        section_embeddings = self._get_embeddings(section_texts)
        matches = {}
        for i, req in enumerate(requirements):
            req_embedding = rfp_embeddings[i:i+1]
            similarities = [
                (self._calculate_similarity(req_embedding, section_embeddings[j:j+1]), section)
                for j, section in enumerate(sections)
            ]
            similarities.sort(key=lambda x: x[0], reverse=True)
            top_matches = [section for score, section in similarities[:3] if score > 0.5]
            for score, section in similarities[:3]:
                if score > 0.5:
                    section.confidence_score = score
            matches[req.section_title] = top_matches
            self.logger.debug(f"Matched {req.section_title} with {len(top_matches)} sections")
        return matches

    def _get_embeddings(self, texts: List[str]) -> List[List[float]]:
        try:
            result = genai.embed_content(
                model=self.embedding_model,
                content=texts,
                task_type="semantic_similarity"
            )
            return result['embedding']
        except Exception as e:
            self.logger.error(f"Error getting embeddings: {e}")
            return [[] for _ in texts]

    def _calculate_similarity(self, emb1: List[float], emb2: List[float]) -> float:
        emb1 = np.array(emb1).flatten()
        emb2 = np.array(emb2).flatten()
        dot_product = np.dot(emb1, emb2)
        norm1 = np.linalg.norm(emb1)
        norm2 = np.linalg.norm(emb2)
        return float(dot_product / (norm1 * norm2)) if norm1 and norm2 else 0.0
    
    def generate_dynamic_section(self, section_type: str, config: Dict, extracted_data: Optional[Dict] = None, requirement: Optional[RFPRequirement] = None, sections: Optional[List[ProposalSection]] = None) -> str:
        """Generate dynamic content for RFP sections using extracted data"""
        
        self.logger.debug(f"Generating content for section_type: {section_type}, config: {config is not None}, extracted_data: {extracted_data is not None}, requirement: {requirement is not None}, sections: {sections is not None}")
        
        # Ensure config is a dictionary
        if not isinstance(config, dict):
            self.logger.error(f"Invalid config: {config}")
            config = {}
        
        if not isinstance(extracted_data, dict):
            self.logger.error(f"Invalid extracted_data for {section_type}: {extracted_data}")
            # extracted_data = {}
            extracted_data = {"text": "", "items": [], "forms": [], "fields": []}
        
        context = f"The RFP is for {config.get('service_type', 'nursing services')} for {config.get('district_name', 'School District')} for {config.get('school_year', '2025/2026')} school year."

        # Initialize section_prompts
        section_prompts = {}

        # Handle requirement section separately
        if section_type == "general_requirements" and requirement:
            proposal_sections = ""
            if sections:
                proposal_sections = "".join([
                    f"{s.section_title}. {s.content[:50]}..."
                    for s in sections[:min(2, len(sections))]
                ])
            section_prompts["general_requirements"] = (
                f"""Generate concise requirements {requirement.section_title} for an RFP for {config.get('service_type')} for {config.get('district_name')}. Use extracted data: {requirement.content[:200]}... Use proposal sections: {proposal_sections}. Define clear requirements in itemized subsection lists, ensuring compliance with {config.get('service_type')} standards. Output within 50 to 70 tokens per list item."""
            )
        else:
            section_prompts = {
            "general_terms_and_conditions": f"""Generate a General Terms And Conditions section for an RFP. Use extracted data: {extracted_data.get('text', '')}. Include objectives (e.g., {config.get('objective', '')}) and background (e.g., {config.get('background', '')}). Describe purpose, scope, submission details, district info, and compliance in Number of paragraphs not in numbered lists formats. Ensure a formal tone, within 200-300 tokens.""",
            
            "scope_of_services": f"""Generate a 'SCOPE OF SERVICES' section for an RFP, providing an overview of required services for {config.get('service_type')}. Use extracted data: {extracted_data.get('text', '')}. Highlight key service requirements from top high-priority needs, focusing on quality and safety. Keep output within 200-300 tokens.""",
            
            "timeline": f"""Generate a timeline section for an RFP, outlining key dates for the RFP process. Use extracted data: {extracted_data.get('text', '')}. Include issue date, submission deadline, and estimated contract award date. Provide the Timeline output within 200-300 tokens.""",
            
            "terms_and_conditions": f"""Generate a terms and conditions section for an RFP. Use extracted data: {', '.join(extracted_data.get('items', []))}. Create the List, Including contract award, rejection rights, payment terms, applicable law, and compliance requirements, formatted as numbered subsections. keep each within 100-150 tokens.""",           
            
            "attachments": f"""Generate an 'ATTACHMENTS' section for an RFP. Use extracted data: {', '.join([f.get('name', '') for f in extracted_data.get('forms', [])] + extracted_data.get('items', []))}. Include all mandatory forms (e.g., Affirmative Action, Non-Collusion) and insurance requirements (e.g., General Liability). List all items concisely, within 200-300 tokens.""",
            
            "appendices": f"""Generate an appendices section for an RFP. Use extracted data: {', '.join(extracted_data.get('items', []))}. List appendices alphabetically as Subsections(e.g., Appendix A, Appendix B, etc) with descriptions. Keep each within 100 to 150 tokens.""",
            
            "submitter_information": f"""Generate a submitter information form section for an RFP. Use extracted data: {', '.join(extracted_data.get('fields', []))}. List required fields (e.g., name, address, phone) for the submitter."""
        }
        # prompt = section_prompts.get(section_type, "Generate a generic RFP section.")
        prompt = f"""You are drafting a professional RFP for a Nursing school district.
        Context: {context}
        Role: Generate a concise section with a formal tone. Use itemized lists for requirements.
        Task: {section_prompts.get(section_type, 'Generate a generic RFP section within specified words')}
        Response:
        """
        try:
            response = self.model.generate_content(
                prompt,
                generation_config={
                    # "max_output_tokens": 150,
                    "temperature": 0.3,
                    "stop_sequences": ["---END---"]
                }
            )
            self.logger.debug(f"Generated {section_type} content: {response.text}")
            return response.text.strip()
        except Exception as e:
            self.logger.error(f"Error generating {section_type} content: {e}")
            # Fallbacks
            if section_type == "general_requirements" and requirement:
                return requirement.content[:200]
            elif section_type == "terms_and_conditions":
                return "\n\n".join(extracted_data.get('items', ["Terms and conditions not specified."]) if extracted_data else ["Terms and conditions not specified."])
            elif section_type == "attachments":
                return "\n".join([f.get('name', 'Form') for f in extracted_data.get('forms', [])] if extracted_data else ["Forms not specified."])
            elif section_type == "appendices":
                return "\n".join(extracted_data.get('items', ["Appendices not specified."]) if extracted_data else ["Appendices not specified."])
            elif section_type == "submitter_information":
                return "\n".join(extracted_data.get('fields', ["Submitter information not specified."]) if extracted_data else ["Submitter information not specified."])
            else:
                return extracted_data.get('text', f"{section_type.title()} not specified.")[:200] if extracted_data else f"{section_type.title()} not specified."

class TemplateManager:
    """Manage RFP template and DOCX generation"""
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def populate_template(self, requirements: List[RFPRequirement], matches: Dict[str, List[ProposalSection]], config: Dict, matcher: ContentMatcher, dynamic_sections: Dict) -> Document:
        if not config:
            self.logger.error("Config is empty or None, cannot populate template")
            raise ValueError("Config dictionary is required to populate template")

        doc = Document()

        # Set document styles
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.line_spacing = WD_LINE_SPACING.ONE_POINT_FIVE

        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Header
        header = sections[0].header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = header_para.add_run(f"{config.get('district_name')}          RFP No. {config.get('rfp_number')}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Footer
        footer = sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Page {PAGE} of {NUMPAGES}"

        # Cover Page
        doc.add_heading(config.get('district_name'), level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('Purchasing Department', style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(config.get('district_address'), style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Phone: {config.get('district_phone')} • Fax: {config.get('district_fax')}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(config.get('district_website'), style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading('REQUEST FOR PROPOSALS', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading(config.get('service_type').upper(), level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"RFP No. {config.get('rfp_number')}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Submission Deadline: {config.get('submission_deadline')}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Issue Date: {config.get('issue_date')}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Prepared by: {config.get('contact_name')}, Purchasing Director", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Contact: {config.get('contact_email')}", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()


        #General terms and conditions
        section_number = 1
        doc.add_heading(f"{section_number}. GENERAL TERMS AND CONDITIONS", level=1)
        intro_content = matcher.generate_dynamic_section(
            section_type="general_terms_and_conditions",
            config=config,
            extracted_data=dynamic_sections.get('introduction', {})
        )
        doc.add_paragraph(intro_content, style='Normal')

        # Scope
        doc.add_heading(f"{section_number}.2 Scope of Services", level=2)
        objective_content = matcher.generate_dynamic_section(
            section_type="scope_of_services",
            config=config,
            extracted_data=dynamic_sections.get('objective', {})
        )
        doc.add_paragraph(objective_content, style='Normal')

        # Project Timeline
        section_number += 1
        doc.add_heading(f"{section_number}. TIMELINE", level=1)
        timeline_content = matcher.generate_dynamic_section(
            section_type="timeline",
            config=config,
            extracted_data=dynamic_sections.get('timeline', {})
        )
        doc.add_paragraph(timeline_content, style='Normal')
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Event"
        table.cell(0, 1).text = "Date"
        table.cell(1, 0).text = "RFP Issue Date"
        table.cell(1, 1).text = config.get('issue_date')
        table.cell(2, 0).text = "Proposal Submission Deadline"
        table.cell(2, 1).text = config.get('submission_deadline')
        table.cell(3, 0).text = "Contract Award (Estimated)"
        table.cell(3, 1).text = (datetime.now() + timedelta(days=60)).strftime('%B %d, %Y')
        
        
        #General requirements
        section_number += 1
        doc.add_heading(f"{section_number}. GENERAL REQUIREMENTS", level=1)
        
        generated_titles = set()
        all_requirements = [req for req in requirements if req.priority >= 0.7][:10] + [req for req in requirements if req.priority < 0.7][:20]  # Limit to 30 requirements
        
        subsection_counter = 1
        for req in all_requirements:
            if req.section_title not in generated_titles:
                # Add subsection heading (level 2)
                # doc.add_heading(f"{section_number}.{subsection_counter} {req.section_title[:80]}", level=2)
                
                # Generate content
                content = matcher.generate_dynamic_section(
                    section_type="general_requirements",
                    config=config,
                    extracted_data={"text": req.content, "items": [req.content[:1000]]},
                    requirement=req,
                    sections=matches.get(req.section_title, [])
                )         
                # Process content and add as bullet points
                if content:
                    content_lines = [line.strip() for line in content.split('\n') if line.strip()]
                    for line in content_lines:
                        # Remove any existing bullet symbols
                        clean_line = line.lstrip('•-* ').strip()
                        if clean_line:
                            bullet_para = doc.add_paragraph(style='Normal')
                            bullet_para.paragraph_format.left_indent = Inches(0.25)
                            bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                            bullet_para.add_run(f"• {clean_line}")
                
                generated_titles.add(req.section_title)
                subsection_counter += 1

        # Terms and Conditions
        section_number += 1
        doc.add_heading(f"{section_number}. Terms and Conditions", level=1)
        terms_content = matcher.generate_dynamic_section(
            section_type="terms_and_conditions",
            config=config,
            extracted_data=dynamic_sections.get('terms_and_conditions', {})
        )
        for idx, term in enumerate(terms_content.split('\n'), 1):
            if term.strip():
                doc.add_heading(f"{section_number}.{idx} {term.split(':')[0]}", level=2)
                doc.add_paragraph(term, style='Normal')
        section_number += 1

        # Mandatory Forms
        doc.add_heading(f"{section_number}. Attachments & Mandatory Forms", level=1)
        forms_content = matcher.generate_dynamic_section(
            section_type="attachments",
            config=config,
            extracted_data=dynamic_sections.get('attachments', {})
        )
        doc.add_paragraph(forms_content, style='Normal')
        forms = dynamic_sections.get('attachments', {}).get('forms', [])[:3]
        for idx, form in enumerate(forms, 1):
            doc.add_heading(f"{section_number}.{idx} {form.get('name', 'Form')}", level=2)
            table = doc.add_table(rows=len(form.get('fields', [])) + 1, cols=1)
            table.style = 'Table Grid'
            table.cell(0, 0).text = form.get('name', 'Form')
            for row, field in enumerate(form.get('fields', []), 1):
                table.cell(row, 0).text = field[:100]
        section_number += 1

        # Appendices
        doc.add_heading(f"{section_number}. Appendices", level=1)
        appendices_content = matcher.generate_dynamic_section(
            section_type="appendices",
            config=config,
            extracted_data=dynamic_sections.get('appendices', {})
        )
        doc.add_paragraph(appendices_content, style='Normal')
        section_number += 1

        # Submitter Information
        doc.add_heading(f"{section_number}. Submitter Information", level=1)
        submitter_content = matcher.generate_dynamic_section(
            section_type="submitter_information",
            config=config,
            extracted_data=dynamic_sections.get('submitter_information', {})
        )
        table = doc.add_table(rows=len(dynamic_sections.get('submitter_information', {}).get('fields', [])[:2]) + 1, cols=1)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Submitter Information"
        for row, field in enumerate(dynamic_sections.get('submitter_information', {}).get('fields', [])[:2], 1):
            table.cell(row, 0).text = field[:100]

        return doc

    
    def save_template(self, doc: Document, doc_file: str):
        try:
            doc.save(doc_file)
            self.logger.info(f"Saved DOCX file: {doc_file}")
        except Exception as e:
            self.logger.error(f"Error saving DOCX file: {e}")
    
    # def save_template(self, doc: Document, doc_file: str):
    #     try:
    #         os.makedirs(os.path.dirname(doc_file) or '.', exist_ok=True)
    #         doc_file = re.sub(r'[<>:"/\\|?*\n\r\t]', '_', doc_file)
    #         # Remove multiple underscores and limit length
    #         doc_file = re.sub(r'_{2,}', '_', doc_file)
    #         if len(doc_file) > 255:
    #             # Keep the extension and truncate the middle part
    #             name, ext = os.path.splitext(doc_file)
    #             doc_file = name[:200] + ext
            
    #         doc.save(doc_file)
    #         absolute_path = os.path.abspath(doc_file)
    #         self.logger.info(f"Saved DOCX file: {absolute_path}")
    #     except Exception as e:
    #         self.logger.error(f"Error saving DOCX file: {e}")

class ValidationChecker:
    """Validate RFP completeness"""
    def __init__(self):
        self.logger = logging.getLogger(__name__)

    def validate(self, requirements: List[RFPRequirement], doc: Document) -> Dict:
        results = {
            'total': len(requirements),
            'addressed': 0,
            'missing': [],
            'mandatory_missing': [],
            'completion_percentage': 0.0
        }
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        for req in requirements:
            if req.section_title in doc_text:
                results['addressed'] += 1
            else:
                results['missing'].append(req.section_title)
                if req.mandatory:
                    results['mandatory_missing'].append(req.section_title)
        results['completion_percentage'] = (results['addressed'] / results['total']) * 100 if results['total'] > 0 else 0.0
        return results

    def generate_report(self, results: Dict) -> str:
        report = f"""
        RFP Validation Report
        =====================
        Completion: {results['completion_percentage']:.1f}%
        Total Sections: {results['total']}
        Addressed: {results['addressed']}
        Missing: {', '.join(results['missing'])}
        Mandatory Missing: {', '.join(results['mandatory_missing'])}
        """
        return report


class RFPGenerator:
    """Generate RFP documents"""
    def __init__(self, gemini_api_key: str):
        genai.configure(api_key=gemini_api_key)
        self.extractor = DocumentExtractor()
        self.parser = RFPParser()
        self.analyzer = ProposalAnalyzer()
        self.matcher = ContentMatcher(gemini_api_key)
        self.templates = TemplateManager()
        self.validator = ValidationChecker()
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def generate_rfp(self, rfp_files: List[str], proposal_files: List[str], output_dir: str) -> Dict:
        os.makedirs(output_dir, exist_ok=True)
        results = {'files_created': []}

        # Validate file extensions and existence
        valid_extensions = ['.pdf', '.docx']
        for file in rfp_files + proposal_files:
            if not os.path.exists(file):
                self.logger.error(f"File not found: {file}")
                raise FileNotFoundError(f"File not found: {file}")
            if os.path.splitext(file)[1].lower() not in valid_extensions:
                self.logger.error(f"Unsupported file format: {file}")
                raise ValueError(f"Unsupported file format: {file}")

        # Extract and parse RFPs
        requirements = []
        rfp_texts = []
        for rfp_file in rfp_files:
            text = self.extractor.extract_text(rfp_file)
            if text:
                rfp_texts.append(text)
                requirements.extend(self.parser.parse_requirements(text))
            else:
                self.logger.warning(f"No text extracted from {rfp_file}")
        results['requirements'] = requirements
        self.logger.info(f"Extracted {len(rfp_texts)} RFP texts and {len(requirements)} requirements")

        # Extract proposals
        proposal_texts = []
        for f in proposal_files:
            text = self.extractor.extract_text(f)
            if text:
                proposal_texts.append(text)
            else:
                self.logger.warning(f"No text extracted from {f}")
        sections = self.analyzer.analyze_proposals(proposal_texts, [os.path.basename(f) for f in proposal_files])
        results['sections'] = sections
        self.logger.info(f"Extracted {len(proposal_texts)} proposal texts and {len(sections)} sections")

        # Generate config dynamically
        gemini_client = genai.GenerativeModel("gemini-1.5-flash")
        config = self.parser.extract_config(rfp_texts[0] if rfp_texts else "", proposal_texts, gemini_client)
        if not config:
            self.logger.error("Failed to generate config")
            raise ValueError("Config generation failed")
        results['config'] = config
        self.logger.info(f"Generated config: {json.dumps(config, indent=2)}")
        
        
        # Extract dynamic sections
        dynamic_sections = self.parser.extract_dynamic_sections(rfp_texts[0] if rfp_texts else "", proposal_texts)
        results['dynamic_sections'] = dynamic_sections
        self.logger.info(f"Extracted dynamic sections: {json.dumps(dynamic_sections, indent=2)}")

        # Match content
        matches = self.matcher.match_requirements(requirements, sections)
        results['matches'] = matches
        self.logger.info(f"Generated {len(matches)} requirement-section matches")

        # Generate DOCX RFP
        doc = self.templates.populate_template(requirements, matches, config, self.matcher, dynamic_sections)
        district_name = config.get('district_name', 'district')
        sanitized_name = re.sub(r'[^\w\s-]', '', district_name)
        sanitized_name = re.sub(r'\s+', '_', sanitized_name.strip())
        sanitized_name = sanitized_name[:50] or 'district'
        # doc_file = os.path.join(output_dir, f"rfp_{config.get('district_name', 'district').replace(' ', '_')}.docx")
        doc_file = os.path.normpath(os.path.join(output_dir, f"rfp_{sanitized_name}.docx"))  # Normalize path
        self.templates.save_template(doc, doc_file)
        results['files_created'].append(doc_file)

        # Validate
        validation = self.validator.validate(requirements, doc)
        report = self.validator.generate_report(validation)
        report_file = os.path.join(output_dir, "validation_report.txt")
        with open(report_file, 'w') as f:
            f.write(report)
        results['files_created'].append(report_file)
        results['validation'] = validation
        self.logger.info(f"Validation results: {json.dumps(validation, indent=2)}")

        return results

# def main():
#     """Run the system"""
#     system = RFPGenerator(gemini_api_key=os.getenv("GOOGLE_API_KEY"))
#     # Define folder paths
#     rfp_folder = "RFP"
#     proposal_folder = "proposal"
#     output_folder = "output"
    
#     # Get all PDF and DOCX files from RFP folder
#     rfp_files = []
#     if os.path.exists(rfp_folder):
#         rfp_files.extend(glob.glob(os.path.join(rfp_folder, "*.pdf")))
#         rfp_files.extend(glob.glob(os.path.join(rfp_folder, "*.docx")))
    
#     proposal_files = []
#     if os.path.exists(proposal_folder):
#         proposal_files.extend(glob.glob(os.path.join(proposal_folder, "*.pdf")))
#         proposal_files.extend(glob.glob(os.path.join(proposal_folder, "*.docx")))
    
#     print(f"Found {len(rfp_files)} RFP files:")
#     for file in rfp_files:
#         print(f"  - {file}")
    
#     print(f"Found {len(proposal_files)} proposal files:")
#     for file in proposal_files:
#         print(f"  - {file}")
    
#     if not rfp_files:
#         print("Error: No RFP files found in 'RFP' folder")
#         return
    
#     if not proposal_files:
#         print("Error: No proposal files found in 'proposal' folder")
#         return
    
#     try:
#         results = system.generate_rfp(rfp_files, proposal_files, output_folder)
#         print(f"Generated files: {results['files_created']}")
#         validation = results.get('validation', {})
#         completion = validation.get('completion_percentage', 0.0)
#         print(f"Completion: {completion:.1f}%")
#     except Exception as e:
#         print(f"Error generating RFP: {e}")

# if __name__ == "__main__":
#     main()